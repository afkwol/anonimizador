"""
Script para crear documento de prueba para el anonimizador
"""
from docx import Document
from docx.shared import Pt
from pathlib import Path

# Crear documento
doc = Document()

# Título
title = doc.add_paragraph()
title_run = title.add_run("SENTENCIA DE PRUEBA")
title_run.bold = True
title_run.font.size = Pt(14)
title.alignment = 1  # Centrado

doc.add_paragraph()

# Contenido de prueba
text = """
JUZGADO CIVIL Y COMERCIAL N° 5
EXPEDIENTE N° 12345/2024

AUTOS: "PÉREZ, JUAN CARLOS c/ GONZÁLEZ, MARÍA LAURA s/ DAÑOS Y PERJUICIOS"

SENTENCIA

En la ciudad de Buenos Aires, a los 15 días del mes de enero de 2025, el Dr. Juan Martínez, Juez del Juzgado Civil y Comercial N° 5, resuelve:

I. ANTECEDENTES

Se presentó el Sr. Juan Carlos Pérez, DNI 12.345.678, CUIL 20-12345678-9, con domicilio en Av. Corrientes 1234, 5° "B", Ciudad Autónoma de Buenos Aires, en contra de la Sra. María Laura González, DNI 98.765.432, domiciliada en Calle Falsa 567, Apto. 8, San Isidro.

El actor reclama la suma de $500.000 por daños materiales y $200.000 por daño moral, derivados de un accidente de tránsito ocurrido el 10 de marzo de 2024.

Según la doctrina de Lorenzetti, "la responsabilidad civil exige la concurrencia de cuatro elementos: antijuridicidad, daño, relación de causalidad y factor de atribución".

II. PRUEBA

Declararon como testigos:
- Pedro López, DNI 11.222.333, quien manifestó haber presenciado el accidente.
- Ana Fernández, empleada del actor, quien declaró sobre las secuelas del accidente.

El perito ingeniero Carlos Rodríguez presentó informe técnico sobre los daños al vehículo.

III. CONSIDERACIONES JURÍDICAS

Como sostiene la Corte Suprema de Justicia de la Nación en Fallos 328:4640, la prueba del daño es carga de quien lo alega.

El fallo de la CSJN en autos "Aquino c/ Cargo Servicios Industriales" estableció importantes precedentes en materia de responsabilidad.

IV. RESOLUTIVO

Por todo lo expuesto, FALLO:

1) Hacer lugar a la demanda interpuesta por Juan Carlos Pérez.
2) Condenar a María Laura González a abonar la suma de $500.000 en concepto de daños materiales.
3) Condenar a la demandada a abonar $200.000 por daño moral.
4) Las costas se imponen a la parte demandada vencida.

Datos de contacto del actor:
- Email: juan.perez@example.com
- Teléfono: 011-4567-8900
- CBU: 0123456789012345678901

FÍRMESE, NOTIFÍQUESE Y ARCHÍVESE.

Dr. Juan Martínez
Juez Civil y Comercial N° 5
"""

# Agregar texto al documento
for paragraph in text.split("\n\n"):
    if paragraph.strip():
        p = doc.add_paragraph(paragraph.strip())
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(11)

# Guardar documento
output_path = Path(__file__).parent / "documento_prueba.docx"
doc.save(str(output_path))

print(f"✓ Documento de prueba creado: {output_path}")
print(f"  Tamaño: {output_path.stat().st_size} bytes")
print()
print("Contenido esperado a anonimizar:")
print("- Actor: Juan Carlos Pérez")
print("- Demandado: María Laura González")
print("- Testigos: Pedro López, Ana Fernández")
print("- Perito: Carlos Rodríguez")
print("- DNIs, CUIL, domicilios, email, teléfono, CBU")
print()
print("Contenido a preservar:")
print("- Juez: Dr. Juan Martínez (magistrado actuante)")
print("- Doctrinarios: Lorenzetti")
print("- Jurisprudencia: CSJN, Fallos 328:4640, Aquino c/ Cargo")
