"""
Procesamiento de documentos judiciales
Soporta: .docx, .doc, .rtf
"""
import platform
import subprocess
from pathlib import Path
from typing import Tuple

from docx import Document
from docx.shared import Pt
from striprtf.striprtf import rtf_to_text


class DocumentProcessingError(Exception):
    """Error durante el procesamiento de documentos"""
    pass


def extract_text(file_path: Path) -> Tuple[str, str]:
    """
    Extrae texto de un documento según su formato

    Args:
        file_path: Ruta al archivo (.docx, .doc, .rtf)

    Returns:
        Tupla (texto_extraido, formato_documento)

    Raises:
        DocumentProcessingError: Si hay error al procesar
        ValueError: Si el formato no es soportado
    """
    if not file_path.exists():
        raise FileNotFoundError(f"Archivo no encontrado: {file_path}")

    ext = file_path.suffix.lower()

    try:
        if ext == ".docx":
            text = _extract_docx(file_path)
            return text, "docx"
        elif ext == ".rtf":
            text = _extract_rtf(file_path)
            return text, "rtf"
        elif ext == ".doc":
            text = _extract_doc(file_path)
            return text, "doc"
        else:
            raise ValueError(
                f"Formato no soportado: {ext}. "
                f"Formatos válidos: .docx, .rtf, .doc"
            )

    except Exception as e:
        if isinstance(e, (ValueError, FileNotFoundError)):
            raise
        raise DocumentProcessingError(
            f"Error al procesar {ext}: {str(e)}"
        ) from e


def _extract_docx(file_path: Path) -> str:
    """
    Extrae texto de archivos .docx

    Args:
        file_path: Ruta al archivo .docx

    Returns:
        Texto extraído con párrafos separados por doble salto

    Raises:
        DocumentProcessingError: Si falla la extracción
    """
    try:
        doc = Document(str(file_path))
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:  # Solo agregar párrafos no vacíos
                paragraphs.append(text)

        if not paragraphs:
            raise DocumentProcessingError("El documento .docx está vacío")

        return "\n\n".join(paragraphs)

    except Exception as e:
        raise DocumentProcessingError(
            f"Error al leer .docx: {str(e)}"
        ) from e


def _extract_rtf(file_path: Path) -> str:
    """
    Extrae texto de archivos .rtf

    Args:
        file_path: Ruta al archivo .rtf

    Returns:
        Texto extraído

    Raises:
        DocumentProcessingError: Si falla la extracción
    """
    try:
        # Intentar diferentes encodings
        encodings = ["utf-8", "latin-1", "cp1252"]

        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    rtf_content = f.read()
                break
            except UnicodeDecodeError:
                continue
        else:
            # Si ninguno funciona, usar errors='ignore'
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                rtf_content = f.read()

        text = rtf_to_text(rtf_content)

        if not text or not text.strip():
            raise DocumentProcessingError("El documento .rtf está vacío")

        return text.strip()

    except Exception as e:
        raise DocumentProcessingError(
            f"Error al leer .rtf: {str(e)}"
        ) from e


def _extract_doc(file_path: Path) -> str:
    """
    Extrae texto de archivos .doc (formato antiguo de Word)

    Args:
        file_path: Ruta al archivo .doc

    Returns:
        Texto extraído

    Raises:
        DocumentProcessingError: Si falla la extracción

    Note:
        - Windows: Requiere Microsoft Word instalado (usa COM automation)
        - Linux/Mac: Requiere antiword instalado
    """
    if platform.system() == "Windows":
        return _extract_doc_windows(file_path)
    else:
        return _extract_doc_antiword(file_path)


def _extract_doc_windows(file_path: Path) -> str:
    """Extrae .doc usando COM automation en Windows"""
    try:
        import win32com.client
    except ImportError:
        raise DocumentProcessingError(
            "Para procesar archivos .doc en Windows necesitas instalar: "
            "pip install pywin32"
        )

    word = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False

        doc = word.Documents.Open(str(file_path.absolute()))
        text = doc.Content.Text
        doc.Close(False)

        if not text or not text.strip():
            raise DocumentProcessingError("El documento .doc está vacío")

        return text.strip()

    except Exception as e:
        raise DocumentProcessingError(
            f"Error al procesar .doc con Word: {str(e)}. "
            f"Asegúrate de que Microsoft Word esté instalado."
        ) from e
    finally:
        if word:
            try:
                word.Quit()
            except:
                pass


def _extract_doc_antiword(file_path: Path) -> str:
    """Extrae .doc usando antiword en Linux/Mac"""
    try:
        result = subprocess.run(
            ["antiword", str(file_path)],
            capture_output=True,
            text=True,
            check=True,
            timeout=30
        )

        text = result.stdout.strip()

        if not text:
            raise DocumentProcessingError("El documento .doc está vacío")

        return text

    except FileNotFoundError:
        raise DocumentProcessingError(
            "antiword no está instalado. Instálalo con:\n"
            "  Ubuntu/Debian: sudo apt-get install antiword\n"
            "  macOS: brew install antiword\n\n"
            "Alternativa: Convierte el archivo a .docx antes de subirlo."
        )
    except subprocess.TimeoutExpired:
        raise DocumentProcessingError(
            "Timeout al procesar .doc. El archivo podría estar corrupto."
        )
    except subprocess.CalledProcessError as e:
        raise DocumentProcessingError(
            f"Error al procesar .doc con antiword: {e.stderr}"
        ) from e


def rebuild_document(original_path: Path, anonymized_text: str) -> Path:
    """
    Reconstruye documento .docx con texto anonimizado

    Args:
        original_path: Ruta al documento original (cualquier formato)
        anonymized_text: Texto anonimizado

    Returns:
        Path al nuevo documento .docx creado

    Note:
        Preserva estructura básica de párrafos pero no formato avanzado
        (negritas, cursivas, tablas, etc.)
    """
    try:
        # Crear nuevo documento
        new_doc = Document()

        # Intentar leer documento original para copiar estilos básicos
        original_ext = original_path.suffix.lower()
        if original_ext == ".docx":
            try:
                original_doc = Document(str(original_path))
                # Copiar estilos del documento original
                if original_doc.styles:
                    new_doc.styles = original_doc.styles
            except:
                pass  # Si falla, usar estilos por defecto

        # Dividir texto en párrafos (doble salto de línea)
        paragraphs = anonymized_text.split("\n\n")

        for para_text in paragraphs:
            if para_text.strip():
                # Agregar párrafo
                p = new_doc.add_paragraph(para_text.strip())

                # Aplicar formato básico
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(11)

        # Generar nombre de archivo de salida
        output_path = original_path.parent / f"anonimizado_{original_path.stem}.docx"

        # Guardar documento
        new_doc.save(str(output_path))

        return output_path

    except Exception as e:
        raise DocumentProcessingError(
            f"Error al reconstruir documento: {str(e)}"
        ) from e


def preserve_document_format(original_path: Path, anonymized_text: str) -> Path:
    """
    Intenta preservar formato del documento original aplicando texto anonimizado

    Args:
        original_path: Ruta al documento original (.docx únicamente)
        anonymized_text: Texto anonimizado

    Returns:
        Path al documento modificado

    Note:
        Solo funciona con .docx. Para otros formatos usa rebuild_document()
    """
    ext = original_path.suffix.lower()

    if ext != ".docx":
        raise ValueError(
            "preserve_document_format solo funciona con .docx. "
            "Usa rebuild_document() para otros formatos."
        )

    try:
        # Cargar documento original
        doc = Document(str(original_path))

        # Dividir texto anonimizado en párrafos
        anon_paragraphs = anonymized_text.split("\n\n")

        # Reemplazar texto de párrafos existentes
        for i, para in enumerate(doc.paragraphs):
            if i < len(anon_paragraphs):
                # Limpiar párrafo actual
                for run in para.runs:
                    run.text = ""

                # Agregar nuevo texto preservando formato del primer run
                if para.runs:
                    para.runs[0].text = anon_paragraphs[i].strip()
                else:
                    para.add_run(anon_paragraphs[i].strip())

        # Si hay más párrafos anonimizados que originales, agregarlos
        for i in range(len(doc.paragraphs), len(anon_paragraphs)):
            doc.add_paragraph(anon_paragraphs[i].strip())

        # Guardar documento modificado
        output_path = original_path.parent / f"anonimizado_{original_path.stem}.docx"
        doc.save(str(output_path))

        return output_path

    except Exception as e:
        raise DocumentProcessingError(
            f"Error al preservar formato: {str(e)}"
        ) from e


# Alias para compatibilidad con código existente
def process_document(file_path: Path) -> Tuple[str, str]:
    """
    Alias de extract_text() para compatibilidad

    Args:
        file_path: Ruta al documento

    Returns:
        Tupla (texto_extraido, formato_documento)
    """
    return extract_text(file_path)
